import socket
import sys
from docx import Document

def scan_ports(host, port_range):
    open_ports = []
    for port in range(*port_range):
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.settimeout(1)
                if not s.connect_ex((host, port)):
                    open_ports.append(port)
        except socket.error:
            pass
    return open_ports

def write_report(host, open_ports, filename):
    doc = Document()
    doc.add_heading(f'Port Scan Report for {host}', 0)

    if open_ports:
        doc.add_paragraph('The following ports are open:')
        for port in open_ports:
            doc.add_paragraph(f'Port {port}')
    else:
        doc.add_paragraph('No open ports found.')

    doc.save(filename)

def process_host(host):
    target_port_range = (20, 1025)
    open_ports = scan_ports(host, target_port_range)
    report_filename = f'port_scan_report_{host}.docx'
    write_report(host, open_ports, report_filename)

# Check if any arguments are provided (for single IP) or read from stdin (for piped list)
if len(sys.argv) > 1:
    process_host(sys.argv[1])
else:
    for line in sys.stdin:
        host = line.strip()
        if host:
            process_host(host)
